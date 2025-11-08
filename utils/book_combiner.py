import os
import shutil
from pypdf import PdfReader, PdfWriter
from docx2pdf import convert as convert_docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import io
from docx import Document
import time
import subprocess
import win32com.client as win32
import os
from win32com.client import constants

def convert_docx_to_pdf_win32(docx_path, pdf_path):
    """
    Convert DOCX to PDF using Microsoft Word with enhanced error handling
    you shall close the word application before running this script.
    """
    word = None
    doc = None
    
    try:
        # Ensure absolute paths
        docx_path = os.path.abspath(docx_path)
        pdf_path = os.path.abspath(pdf_path)
        
        # Check if source file exists
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Source file not found: {docx_path}")
        
        # Create directory for PDF if it doesn't exist
        pdf_dir = os.path.dirname(pdf_path)
        if pdf_dir and not os.path.exists(pdf_dir):
            os.makedirs(pdf_dir)
        
        word = win32.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0  # Disable alerts
        
        print(f"Opening document: {docx_path}")
        doc = word.Documents.Open(docx_path)
        
        print(f"Saving as PDF: {pdf_path}")
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
        print("Conversion successful")
        
    except Exception as e:
        error_msg = str(e)
        print(f"Error details: {error_msg}")
        raise Exception(f"Failed to convert {docx_path} to PDF: {error_msg}")
    
    finally:
        # Clean up
        if doc:
            try:
                doc.Close(SaveChanges=0)  # Don't save changes
            except:
                pass
        if word:
            try:
                word.Quit()
            except:
                pass

def validate_docx_file(file_path):
    """
    Validate DOCX file before conversion
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File does not exist: {file_path}")
    
    if os.path.getsize(file_path) == 0:
        raise ValueError(f"File is empty: {file_path}")
    
    # Try to open with python-docx to check if it's valid
    try:
        from docx import Document
        doc = Document(file_path)
        # Just try to access basic properties
        _ = doc.paragraphs[0].text if doc.paragraphs else ""
    except Exception as e:
        raise ValueError(f"Invalid DOCX file: {file_path}, Error: {e}")

def convert_docx_to_pdf_pandoc(docx_path, pdf_path):
    """
    Convert DOCX to PDF using Pandoc  --> the format is not good.
    """
    try:
        # Ensure absolute paths
        docx_path = os.path.abspath(docx_path)
        pdf_path = os.path.abspath(pdf_path)
        
        # Create directory for PDF if it doesn't exist
        pdf_dir = os.path.dirname(pdf_path)
        if pdf_dir and not os.path.exists(pdf_dir):
            os.makedirs(pdf_dir)
        
        # Build pandoc command
        cmd = [
            'D:\Programs\pandoc-3.5\pandoc',
            docx_path,
            '-o', pdf_path,
            '--pdf-engine=xelatex',  # Better handling of complex layouts
            '--variable', 'geometry:margin=1in'  # Standard margins
        ]
        
        print(f"Converting {docx_path} to {pdf_path} using pandoc...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode != 0:
            raise Exception(f"Pandoc conversion failed: {result.stderr}")
        
        if not os.path.exists(pdf_path):
            raise Exception("PDF file was not created")
            
        print("Pandoc conversion successful")
        return pdf_path
        
    except subprocess.TimeoutExpired:
        raise Exception("Pandoc conversion timed out")
    except FileNotFoundError:
        raise Exception("Pandoc not found. Please install pandoc from https://pandoc.org/")
    except Exception as e:
        raise Exception(f"Pandoc conversion error: {str(e)}")

# def create_pdf_book_with_toc(docx_files, topics, output_pdf_path, conversion='win32'):
#     """
#     Converts, combines, and paginates a list of DOCX files into a single PDF book
#     with a generated table of contents.

#     Args:
#         docx_files (list): List of paths to the .docx files.
#         topics (list): List of chapter topic strings.
#         output_pdf_path (str): The full path for the final output PDF.
#     """
#     temp_dir = "_temp_pdf_processing"
#     if os.path.exists(temp_dir):
#         shutil.rmtree(temp_dir)
#     os.makedirs(temp_dir)

#     print("--- Step 1: Converting DOCX to PDF and counting pages ---")
#     chapter_data = []
#     for docx_file in docx_files:
#         try:
#             pdf_path = os.path.join(temp_dir, os.path.basename(docx_file).replace('.docx', '.pdf'))
#             print(f"Converting {docx_file} to {pdf_path}...")
#             validate_docx_file(docx_file)
#             if conversion == 'win32':
#                 convert_docx_to_pdf_win32(docx_file, pdf_path)
#             elif conversion == 'pandoc':
#                 convert_docx_to_pdf_pandoc(docx_file, pdf_path)
#             else:
#                 convert_docx(docx_file, pdf_path)
#             time.sleep(2)

#             reader = PdfReader(pdf_path)
#             page_count = len(reader.pages)
#             chapter_data.append({'path': pdf_path, 'pages': page_count})
#             print(f" -> Found {page_count} pages.")
#         except Exception as e:
#             print(f"Could not process {docx_file}. Skipping. Error: {e}")

#     print("\n--- Step 2: Calculating Table of Contents ---")
#     toc_entries = []
#     # The TOC itself is page 1, so the first chapter starts on page 2.
#     current_page_number = 2
#     for i, data in enumerate(chapter_data):
#         toc_entries.append({'topic': topics[i], 'page': current_page_number})
#         current_page_number += data['pages']

#     print("TOC calculation complete. total pages:", current_page_number, "pages")
#     for entry in toc_entries:
#         print(f" -> {entry['topic']}: Page {entry['page']}")

#     print("\n--- Step 3: Creating the Contents page PDF ---")
#     toc_pdf_path = os.path.join(temp_dir, "toc.pdf")
#     c = canvas.Canvas(toc_pdf_path, pagesize=letter)
#     width, height = letter
#     c.setFont("Helvetica-Bold", 18)
#     c.drawCentredString(width / 2.0, height - 1.5 * inch, "Contents")
#     c.setFont("Helvetica", 12)
#     text_object = c.beginText(1.5 * inch, height - 2.5 * inch)
#     for entry in toc_entries:
#         text_object.textLine(f"{entry['topic']} .................... {entry['page']}")
#     c.drawText(text_object)
#     c.save()
#     print(f"Contents page saved to {toc_pdf_path}")

#     print("\n--- Step 4: Merging all chapter PDFs ---")
#     merged_chapters_path = os.path.join(temp_dir, "merged_chapters.pdf")
#     merger = PdfWriter()
#     for data in chapter_data:
#         merger.append(data['path'])
#     merger.write(merged_chapters_path)
#     merger.close()
#     print(f"All chapters merged into {merged_chapters_path}")

#     print("\n--- Step 5: Assembling final book (TOC + Chapters) ---")
#     final_book_no_pages_path = os.path.join(temp_dir, "final_book_no_pages.pdf")
#     final_merger = PdfWriter()
#     final_merger.append(toc_pdf_path)
#     final_merger.append(merged_chapters_path)
#     final_merger.write(final_book_no_pages_path)
#     final_merger.close()
#     print("Final book assembled.")

#     print("\n--- Step 6: Adding page numbers to the final book ---")
#     reader = PdfReader(final_book_no_pages_path)
#     writer = PdfWriter()
#     for i, page in enumerate(reader.pages):
#         packet = io.BytesIO()
#         # Create a new PDF with just a page number
#         c = canvas.Canvas(packet, pagesize=letter)
#         c.setFont("Helvetica", 9)
#         # Position the page number in the bottom right
#         c.drawString(width - 1 * inch, 0.75 * inch, str(i + 1))
#         c.save()

#         # Move to the beginning of the StringIO buffer
#         packet.seek(0)

#         # Create a new PDF from the page number packet
#         overlay = PdfReader(packet).pages[0]
#         page.merge_page(overlay)
#         writer.add_page(page)

#     with open(output_pdf_path, "wb") as f:
#         writer.write(f)
#     print(f"Page numbers added. Final book saved to: {output_pdf_path}")

#     print("\n--- Step 7: Cleaning up temporary files ---")
#     shutil.rmtree(temp_dir)
#     print("Cleanup complete.")

# def create_pdf_book_with_toc(docx_files, topics, output_pdf_path, conversion='win32'):
#     """
#     Converts, combines, and paginates a list of DOCX files into a single PDF book
#     with a generated table of contents.

#     Args:
#         docx_files (list): List of paths to the .docx files.
#         topics (list): List of chapter topic strings.
#         output_pdf_path (str): The full path for the final output PDF.
#     """
#     temp_dir = "_temp_pdf_processing"
#     if os.path.exists(temp_dir):
#         shutil.rmtree(temp_dir)
#     os.makedirs(temp_dir)

#     print("--- Step 1: Converting DOCX to PDF and counting pages ---")
#     chapter_data = []
#     for docx_file in docx_files:
#         try:
#             pdf_path = os.path.join(temp_dir, os.path.basename(docx_file).replace('.docx', '.pdf'))
#             print(f"Converting {docx_file} to {pdf_path}...")
#             validate_docx_file(docx_file)
#             if conversion == 'win32':
#                 convert_docx_to_pdf_win32(docx_file, pdf_path)
#             elif conversion == 'pandoc':
#                 convert_docx_to_pdf_pandoc(docx_file, pdf_path)
#             else:
#                 convert_docx(docx_file, pdf_path)
#             time.sleep(2)

#             reader = PdfReader(pdf_path)
#             page_count = len(reader.pages)
#             chapter_data.append({'path': pdf_path, 'pages': page_count})
#             print(f" -> Found {page_count} pages.")
#         except Exception as e:
#             print(f"Could not process {docx_file}. Skipping. Error: {e}")

#     print("\n--- Step 2: Calculating Table of Contents ---")
#     toc_entries = []
#     # The TOC itself is page 1, so the first chapter starts on page 2.
#     current_page_number = 2
#     for i, data in enumerate(chapter_data):
#         toc_entries.append({'topic': topics[i], 'page': current_page_number})
#         current_page_number += data['pages']

#     print("TOC calculation complete. total pages:", current_page_number, "pages")
#     for entry in toc_entries:
#         print(f" -> {entry['topic']}: Page {entry['page']}")

#     print("\n--- Step 3: Creating the Contents page PDF ---")
#     toc_pdf_path = os.path.join(temp_dir, "toc.pdf")
#     c = canvas.Canvas(toc_pdf_path, pagesize=letter)
#     width, height = letter
#     c.setFont("Helvetica-Bold", 18)
#     c.drawCentredString(width / 2.0, height - 1.5 * inch, "Contents")
#     c.setFont("Helvetica", 12)
#     text_object = c.beginText(1.5 * inch, height - 2.5 * inch)
#     for entry in toc_entries:
#         text_object.textLine(f"{entry['topic']} .................... {entry['page']}")
#     c.drawText(text_object)
#     c.save()
#     print(f"Contents page saved to {toc_pdf_path}")

#     print("\n--- Step 4: Merging all chapter PDFs ---")
#     merged_chapters_path = os.path.join(temp_dir, "merged_chapters.pdf")
#     merger = PdfWriter()
    
#     # Add verification for each chapter before merging
#     total_expected_pages = 0
#     for i, data in enumerate(chapter_data):
#         try:
#             reader = PdfReader(data['path'])
#             page_count = len(reader.pages)
#             print(f"Adding chapter {i+1} ({data['path']}): {page_count} pages")
#             merger.append(data['path'])
#             total_expected_pages += page_count
#         except Exception as e:
#             print(f"Error adding chapter {i+1} from {data['path']}: {e}")
    
#     merger.write(merged_chapters_path)
#     merger.close()
    
#     # Verify merged chapters
#     chapters_reader = PdfReader(merged_chapters_path)
#     actual_pages = len(chapters_reader.pages)
#     print(f"All chapters merged into {merged_chapters_path}")
#     print(f"Expected pages: {total_expected_pages}, Actual pages: {actual_pages}")
#     if total_expected_pages != actual_pages:
#         print("WARNING: Page count mismatch in merged chapters!")

#     print("\n--- Step 5: Assembling final book (TOC + Chapters) ---")
#     final_book_no_pages_path = os.path.join(temp_dir, "final_book_no_pages.pdf")
#     final_merger = PdfWriter()
    
#     # Add TOC
#     toc_reader = PdfReader(toc_pdf_path)
#     print(f"Adding TOC: {len(toc_reader.pages)} pages")
#     final_merger.append(toc_pdf_path)
    
#     # Add chapters
#     final_merger.append(merged_chapters_path)
#     final_merger.write(final_book_no_pages_path)
#     final_merger.close()
    
#     # Verify final assembly
#     final_reader = PdfReader(final_book_no_pages_path)
#     expected_final_pages = len(toc_reader.pages) + len(chapters_reader.pages)
#     actual_final_pages = len(final_reader.pages)
#     print(f"Final book assembled: {actual_final_pages} pages")
#     print(f"Expected pages: {expected_final_pages}, Actual pages: {actual_final_pages}")
#     if expected_final_pages != actual_final_pages:
#         print("WARNING: Page count mismatch in final assembly!")
    
#     print("\n--- Step 6: Adding page numbers to the final book ---")
#     reader = PdfReader(final_book_no_pages_path)
#     writer = PdfWriter()
    
#     # Get page dimensions from the first page
#     if reader.pages:
#         page_box = reader.pages[0].mediabox
#         page_width = page_box.width
#         page_height = page_box.height
#     else:
#         page_width, page_height = letter
    
#     print(f"Adding page numbers to {len(reader.pages)} pages...")
#     for i, page in enumerate(reader.pages):
#         try:
#             packet = io.BytesIO()
#             # Create a new PDF with just a page number
#             c = canvas.Canvas(packet, pagesize=(page_width, page_height))
#             c.setFont("Helvetica", 9)
#             # Position the page number in the bottom right
#             c.drawString(page_width - 1 * inch, 0.75 * inch, str(i + 1))
#             c.save()

#             # Move to the beginning of the buffer
#             packet.seek(0)

#             # Create a new PDF from the page number packet
#             overlay_reader = PdfReader(packet)
#             if overlay_reader.pages:
#                 overlay = overlay_reader.pages[0]
#                 page.merge_page(overlay)
#             else:
#                 print(f"Warning: Could not create page number overlay for page {i+1}")
            
#             writer.add_page(page)
#         except Exception as e:
#             print(f"Error processing page {i+1}: {e}")
#             # Add the page without page number if there's an error
#             writer.add_page(page)

#     with open(output_pdf_path, "wb") as f:
#         writer.write(f)
#     print(f"Page numbers added. Final book saved to: {output_pdf_path}")

#     print("\n--- Step 7: Cleaning up temporary files ---")
#     shutil.rmtree(temp_dir)
#     print("Cleanup complete.")

def create_pdf_book_with_toc(docx_files, topics, output_pdf_path, conversion='win32', cover_file_docx=None):
    """
    Converts, combines, and paginates a list of DOCX files into a single PDF book
    with a generated table of contents.

    Args:
        docx_files (list): List of paths to the .docx files.
        topics (list): List of chapter topic strings.
        output_pdf_path (str): The full path for the final output PDF.
        conversion (str): Conversion method ('win32', 'pandoc', or 'docx2pdf').
        cover_file_docx (str): Path to the cover page DOCX file (optional).
    """
    temp_dir = "_temp_pdf_processing"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    print("--- Step 1: Converting DOCX to PDF and counting pages ---")
    chapter_data = []
    
    # Process cover page if provided
    cover_pdf_path = None
    if cover_file_docx:
        # Handle case where cover_file_docx might be a tuple (e.g., from trailing comma)
        if isinstance(cover_file_docx, tuple):
            cover_file_docx = cover_file_docx[0] if cover_file_docx else None
            
        if cover_file_docx and os.path.exists(cover_file_docx):
            try:
                cover_pdf_path = os.path.join(temp_dir, "cover.pdf")
                print(f"Converting cover page {cover_file_docx} to {cover_pdf_path}...")
                validate_docx_file(cover_file_docx)
                if conversion == 'win32':
                    convert_docx_to_pdf_win32(cover_file_docx, cover_pdf_path)
                elif conversion == 'pandoc':
                    convert_docx_to_pdf_pandoc(cover_file_docx, cover_pdf_path)
                else:
                    convert_docx(cover_file_docx, cover_pdf_path)
                time.sleep(2)
                print("Cover page converted successfully.")
            except Exception as e:
                print(f"Could not process cover page {cover_file_docx}. Error: {e}")
                cover_pdf_path = None
        elif cover_file_docx:
            print(f"Cover file {cover_file_docx} does not exist.")
    
    for docx_file in docx_files:
        try:
            pdf_path = os.path.join(temp_dir, os.path.basename(docx_file).replace('.docx', '.pdf'))
            print(f"Converting {docx_file} to {pdf_path}...")
            validate_docx_file(docx_file)
            if conversion == 'win32':
                convert_docx_to_pdf_win32(docx_file, pdf_path)
            elif conversion == 'pandoc':
                convert_docx_to_pdf_pandoc(docx_file, pdf_path)
            else:
                convert_docx(docx_file, pdf_path)
            time.sleep(2)

            reader = PdfReader(pdf_path)
            page_count = len(reader.pages)
            chapter_data.append({'path': pdf_path, 'pages': page_count})
            print(f" -> Found {page_count} pages.")
        except Exception as e:
            print(f"Could not process {docx_file}. Skipping. Error: {e}")

    print("\n--- Step 2: Calculating Table of Contents ---")
    toc_entries = []
    # The TOC itself is page 1, so the first chapter starts on page 2.
    # If we have a cover page, first chapter starts on page 3.
    current_page_number = 2 if not cover_pdf_path else 3
    for i, data in enumerate(chapter_data):
        toc_entries.append({'topic': topics[i], 'page': current_page_number})
        current_page_number += data['pages']

    print("TOC calculation complete. total pages:", current_page_number, "pages")
    for entry in toc_entries:
        print(f" -> {entry['topic']}: Page {entry['page']}")

    print("\n--- Step 3: Creating the Contents page PDF ---")
    toc_pdf_path = os.path.join(temp_dir, "toc.pdf")
    c = canvas.Canvas(toc_pdf_path, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2.0, height - 1.5 * inch, "Contents")
    c.setFont("Helvetica", 12)
    text_object = c.beginText(1.5 * inch, height - 2.5 * inch)
    for entry in toc_entries:
        text_object.textLine(f"{entry['topic']} .................... {entry['page']}")
    c.drawText(text_object)
    c.save()
    print(f"Contents page saved to {toc_pdf_path}")

    print("\n--- Step 4: Merging all chapter PDFs ---")
    merged_chapters_path = os.path.join(temp_dir, "merged_chapters.pdf")
    merger = PdfWriter()
    
    # Add verification for each chapter before merging
    total_expected_pages = 0
    for i, data in enumerate(chapter_data):
        try:
            reader = PdfReader(data['path'])
            page_count = len(reader.pages)
            print(f"Adding chapter {i+1} ({data['path']}): {page_count} pages")
            merger.append(data['path'])
            total_expected_pages += page_count
        except Exception as e:
            print(f"Error adding chapter {i+1} from {data['path']}: {e}")
    
    merger.write(merged_chapters_path)
    merger.close()
    
    # Verify merged chapters
    chapters_reader = PdfReader(merged_chapters_path)
    actual_pages = len(chapters_reader.pages)
    print(f"All chapters merged into {merged_chapters_path}")
    print(f"Expected pages: {total_expected_pages}, Actual pages: {actual_pages}")
    if total_expected_pages != actual_pages:
        print("WARNING: Page count mismatch in merged chapters!")

    print("\n--- Step 5: Assembling final book ---")
    final_book_no_pages_path = os.path.join(temp_dir, "final_book_no_pages.pdf")
    final_merger = PdfWriter()
    
    # Add cover page if available
    if cover_pdf_path:
        try:
            cover_reader = PdfReader(cover_pdf_path)
            print(f"Adding cover page: {len(cover_reader.pages)} pages")
            final_merger.append(cover_pdf_path)
        except Exception as e:
            print(f"Error adding cover page: {e}")
    
    # Add TOC
    toc_reader = PdfReader(toc_pdf_path)
    print(f"Adding TOC: {len(toc_reader.pages)} pages")
    final_merger.append(toc_pdf_path)
    
    # Add chapters
    final_merger.append(merged_chapters_path)
    final_merger.write(final_book_no_pages_path)
    final_merger.close()
    
    # Verify final assembly
    final_reader = PdfReader(final_book_no_pages_path)
    expected_final_pages = len(chapters_reader.pages) + len(toc_reader.pages)
    if cover_pdf_path:
        try:
            cover_reader = PdfReader(cover_pdf_path)
            expected_final_pages += len(cover_reader.pages)
        except Exception as e:
            print(f"Error reading cover page: {e}")
    
    actual_final_pages = len(final_reader.pages)
    print(f"Final book assembled: {actual_final_pages} pages")
    print(f"Expected pages: {expected_final_pages}, Actual pages: {actual_final_pages}")
    if expected_final_pages != actual_final_pages:
        print("WARNING: Page count mismatch in final assembly!")
    
    print("\n--- Step 6: Adding page numbers to the final book ---")
    reader = PdfReader(final_book_no_pages_path)
    writer = PdfWriter()
    
    # Get page dimensions from the first page
    if reader.pages:
        page_box = reader.pages[0].mediabox
        page_width = float(page_box.width)
        page_height = float(page_box.height)
    else:
        page_width, page_height = letter
    
    print(f"Adding page numbers to {len(reader.pages)} pages...")
    cover_page_count = 0
    if cover_pdf_path:
        try:
            cover_reader = PdfReader(cover_pdf_path)
            cover_page_count = len(cover_reader.pages)
        except Exception as e:
            print(f"Error reading cover page for page numbering: {e}")
            
    for i, page in enumerate(reader.pages):
        # Skip page numbering for cover page (if exists) and TOC
        if i < cover_page_count:
            # No page number for cover page
            writer.add_page(page)
            continue
        elif i < cover_page_count + 1:  # TOC page
            # No page number for TOC
            writer.add_page(page)
            continue
        
        try:
            # Create a new PDF with just a page number
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=(page_width, page_height))
            c.setFont("Helvetica", 9)
            # Position the page number in the bottom right
            c.drawString(page_width - 1 * inch, 0.75 * inch, str(i + 1 - cover_page_count))
            c.save()
            packet.seek(0)
            
            # Create overlay PDF
            overlay_reader = PdfReader(packet)
            
            # Create a new page object to avoid modifying the original
            new_page = writer.add_blank_page(width=page_width, height=page_height)
            
            # Merge content from original page
            if overlay_reader.pages:
                new_page.merge_page(page)  # Original content first
                new_page.merge_page(overlay_reader.pages[0])  # Then page number
            else:
                new_page.merge_page(page)  # Only original content if overlay failed
                
        except Exception as e:
            print(f"Error processing page {i+1}: {e}")
            # Add the original page without page number if there's an error
            writer.add_page(page)

    with open(output_pdf_path, "wb") as f:
        writer.write(f)
    print(f"Page numbers added. Final book saved to: {output_pdf_path}")

    print("\n--- Step 7: Cleaning up temporary files ---")
    # shutil.rmtree(temp_dir)
    print("Cleanup complete.")

# --- Example Usage ---
if __name__ == '__main__':
    # IMPORTANT: Create these dummy files or replace with your actual file paths.
    chapter_files_docx = [
        r'D:\Writing\RAG\Ch0 preface.docx', 
        r'D:\Writing\RAG\Ch1 Introduction.docx', 
        r'D:\Writing\RAG\Ch2 LLMOps.docx', 
        r'D:\Writing\RAG\Ch3 RAG.docx',
        r'D:\Writing\RAG\Ch4 data process.docx', 
        r'D:\Writing\RAG\Ch5 Vector.docx',
        r'D:\Writing\RAG\Ch6 Query.docx',
        r'D:\Writing\RAG\Ch7 retrieval.docx',
        r'D:\Writing\RAG\Ch8 augumentation.docx',
        r'D:\Writing\RAG\Ch9 generation.docx',
        r'D:\Writing\RAG\Ch10 Evaluation.docx',
        r'D:\Writing\RAG\Ch11 Serving Monitoring.docx',
        r'D:\Writing\RAG\Ch12 pipeline.docx', 
        r'D:\Writing\RAG\Ch13 NL2SQL.docx',
        r'D:\Writing\RAG\Ch14 GraphRAG.docx',  
        r'D:\Writing\RAG\Ch15 Agentic RAG.docx',   
        r'D:\Writing\RAG\Ch16 Conclusion.docx',
        r'D:\Writing\RAG\Ch17 Appendix.docx',    
    ]

    cover_file_docx = r'D:\Writing\RAG\cover.docx',    
    # Corresponding topics for the table of contents
    chapter_titles = [
        "Preface",  #0
        "1. Introduction",  #1
        "2. MLOps, LLMOps and RAGOps for Production", #2
        "3. RAG Challenges and Solutions", #3
        "4. Data Processing", #4
        "5. Embedding and Vector Database", # 5
        "6. Query Transformation and Prompt Engineering", #6
        "7. Retrieval Techniques", #7 
        "8. Augmentation and Refinement Techniques", #8 
        "9. Generation Techniques", #9
        "10. Evaluation Methodology", # 10
        "11. Serving and Monitoring", # 11
        "12. Pipeline and Orchestration", # 12
        "13. RAG with Database and Text2SQL", #13
        "14. GraphRAG", #14
        "15. Agentic RAG", #15
        "16. Conclusion", #16
        "Appendix" #17
    ]

    # Create dummy docx files for the script to run
    print("Creating dummy .docx files for demonstration...")
    for i, filename in enumerate(chapter_files_docx):
        if not os.path.exists(filename):
            doc = Document()
            doc.add_heading(chapter_titles[i], level=1)
            doc.add_paragraph(f"This is the main content for {chapter_titles[i]}.")
            # Add extra paragraphs to simulate multiple pages
            for p in range(100): # This should create at least 2 pages
                doc.add_paragraph(f"This is some filler text for paragraph {p+1} to ensure the document has content.")
            doc.save(filename)

    # Define the final output file name
    final_output_file = "Complete_Book.pdf"

    # Run the main function
    create_pdf_book_with_toc(chapter_files_docx, chapter_titles, final_output_file, cover_file_docx=cover_file_docx, conversion='win32')